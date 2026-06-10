#!/usr/bin/env python3
"""
Alignment runner for SRT Workbench macOS app.

Reads a JSON config from stdin, runs CTC forced alignment, and outputs
progress + results as JSON lines to stdout.

Input JSON:
{
    "audio_path": "/path/to/audio.wav",
    "script_lines": ["Line one", "Line two", ...],
    "output_path": "/path/to/output.srt",
    "model_path": "/path/to/model.onnx",
    "mode": "align" | "extract"      # extract = parse .docx only, no alignment
}

Output JSON lines:
{"type": "progress", "stage": "...", "percent": 0-100}
{"type": "info", "words": N, "audio_seconds": S}
{"type": "section_match", "matched": bool, "heading": "..."}
{"type": "extract_result", "heading": "...", "lines": [...]}
{"type": "result", "srt_path": "...", "num_cues": N}
{"type": "error", "message": "..."}

All diagnostic logging goes to stderr (captured by the Swift app's log file).
"""

import json
import math
import os
import re
import sys
import threading
import time
from datetime import datetime


def log(message):
    """Timestamped diagnostic line on stderr — captured by the app's log."""
    ts = datetime.now().strftime("%H:%M:%S.%f")[:-3]
    print(f"[{ts}] {message}", file=sys.stderr, flush=True)


def log_environment():
    """Log interpreter + package versions for remote diagnostics."""
    log(f"python {sys.version.split()[0]} ({sys.executable})")
    try:
        from importlib import metadata
        for pkg in ("ctc-forced-aligner", "python-docx", "onnxruntime", "Unidecode"):
            try:
                log(f"package {pkg}=={metadata.version(pkg)}")
            except metadata.PackageNotFoundError:
                log(f"package {pkg}: not installed")
    except Exception as e:  # noqa: BLE001 — diagnostics must never kill the run
        log(f"could not read package versions: {e}")


def progress(stage, percent=0):
    """Send a progress update to the Swift app via stdout."""
    msg = {"type": "progress", "stage": stage, "percent": percent}
    print(json.dumps(msg), flush=True)


def info(words, audio_seconds):
    """Send run context (script size, audio length) to the Swift app."""
    msg = {"type": "info", "words": words, "audio_seconds": audio_seconds}
    print(json.dumps(msg), flush=True)


def error(message):
    """Send an error to the Swift app via stdout."""
    log(f"ERROR: {message}")
    msg = {"type": "error", "message": message}
    print(json.dumps(msg), flush=True)
    sys.exit(1)


def result(srt_path, num_cues):
    """Send the final result to the Swift app via stdout."""
    msg = {"type": "result", "srt_path": srt_path, "num_cues": num_cues}
    print(json.dumps(msg), flush=True)


def section_match(matched, heading):
    """Report whether a document section was matched to the video."""
    msg = {"type": "section_match", "matched": matched, "heading": heading}
    print(json.dumps(msg), flush=True)


def extract_result(heading, lines):
    """Send the extracted script lines (extract mode) to the Swift app."""
    msg = {"type": "extract_result", "heading": heading, "lines": lines}
    print(json.dumps(msg), flush=True)


def format_srt_time(seconds):
    """Convert seconds to SRT timestamp format HH:MM:SS,mmm.

    Rounds to whole milliseconds FIRST so a value like 59.9996 becomes
    00:01:00,000 rather than the invalid 00:00:59,1000.
    """
    total_ms = max(0, int(round(seconds * 1000)))
    h, rem = divmod(total_ms, 3_600_000)
    m, rem = divmod(rem, 60_000)
    s, ms = divmod(rem, 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _compile_filters(filter_patterns):
    """Compile a list of regex pattern strings, reporting invalid ones as errors."""
    compiled = []
    for p in (filter_patterns or []):
        try:
            compiled.append(re.compile(p))
        except re.error as e:
            error(f"Invalid filter regex '{p}': {e}")
    return compiled


def _is_spoken_line(text, filters=None):
    """Return True if the line is non-empty spoken text that doesn't match any filter."""
    if not text:
        return False
    if filters:
        for pattern in filters:
            if pattern.match(text):
                return False
    return True


def _apply_strip_patterns(text, strips):
    """Remove all inline matches of strip patterns from text."""
    for pattern in strips:
        text = pattern.sub("", text)
    # Collapse multiple spaces and re-strip
    text = re.sub(r"  +", " ", text).strip()
    return text


def extract_script_lines(docx_path, filter_patterns=None, strip_patterns=None):
    """Extract spoken lines from a .docx script file using python-docx.

    First strips inline patterns (e.g. slide numbers), then filters out
    empty paragraphs and lines matching any of the provided filter patterns.
    """
    from docx import Document

    filters = _compile_filters(filter_patterns)
    strips = _compile_filters(strip_patterns)
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if strips:
            text = _apply_strip_patterns(text, strips)
        if _is_spoken_line(text, filters):
            lines.append(text)
    return lines


def _normalize_identifier(s):
    """Normalize a string for fuzzy matching: lowercase, collapse separators."""
    s = s.lower()
    s = re.sub(r"[._\-\s]+", " ", s)
    return s.strip()


def _extract_leading_id(s):
    """Extract leading identifier like '1.2.1', '3a', 'Ch3' from a string."""
    m = re.match(r"^\s*([A-Za-z]*\d[\d.]*[A-Za-z]?)", s)
    return m.group(1).lower() if m else None


def _find_matching_section(sections, video_stem):
    """Find the section whose heading matches the video filename stem.

    Returns (heading_text, lines) tuple or (None, None) if no match.
    """
    stem_norm = _normalize_identifier(video_stem)

    # Strategy 1: normalized stem is a substring of heading or vice versa
    for heading, lines in sections:
        heading_norm = _normalize_identifier(heading)
        if stem_norm in heading_norm or heading_norm in stem_norm:
            return heading, lines

    # Strategy 2: compare leading numeric/alphanumeric identifiers
    stem_id = _extract_leading_id(video_stem)
    if stem_id:
        for heading, lines in sections:
            heading_id = _extract_leading_id(heading)
            if heading_id and stem_id == heading_id:
                return heading, lines

    return None, None


def _is_section_heading(style_name, text):
    """Detect whether a paragraph is a section heading.

    Checks Word heading styles first, then falls back to detecting text
    patterns like '1.2.1 – Title', 'Ch3 - Title', 'A2: Title', etc.
    """
    if style_name.startswith("Heading"):
        return True
    # Text pattern: dotted identifier like '1.2.1' or letter-prefixed like 'Ch3'
    # followed by a separator (dash, en-dash, em-dash, colon, pipe) and a title.
    # Requires either a dot in the identifier or a letter prefix to avoid
    # matching slide numbers like '1:' or '2:'.
    if re.match(r"^[A-Za-z]+\d[\d.]*[A-Za-z]?\s*[–\-—:|]", text):
        return True
    if re.match(r"^[A-Za-z]*\d[\d.]*[A-Za-z]?\s*[–\-—:|]", text) and "." in text.split()[0]:
        return True
    return False


def extract_section_lines(docx_path, video_stem, filter_patterns=None, strip_patterns=None):
    """Extract spoken lines for a specific section of a multi-section .docx.

    Splits the document on section headings (detected by Word heading styles
    or text patterns like '1.2.1 – Title'), then matches the video filename
    stem against heading text. Returns (heading, lines) if matched, or
    (None, None) if no headings exist or no match is found.
    """
    from docx import Document

    filters = _compile_filters(filter_patterns)
    strips = _compile_filters(strip_patterns)
    doc = Document(docx_path)
    sections = []  # list of (heading_text, [lines])
    current_heading = None
    current_lines = []

    for para in doc.paragraphs:
        style_name = para.style.name or ""
        text = para.text.strip()

        if text and _is_section_heading(style_name, text):
            # Save previous section
            if current_heading is not None:
                sections.append((current_heading, current_lines))
            current_heading = text
            current_lines = []
        else:
            if strips:
                text = _apply_strip_patterns(text, strips)
            if _is_spoken_line(text, filters):
                current_lines.append(text)

    # Capture the last section
    if current_heading is not None:
        sections.append((current_heading, current_lines))

    if not sections:
        return None, None

    return _find_matching_section(sections, video_stem)


def _split_line_segments(line_words, line_word_ts, max_secs, max_chars):
    """Split a line's words into segments respecting both a max duration and a
    max character count, breaking only at word boundaries. Uses the real
    word-level timestamps so each segment keeps accurate timing.

    Returns a list of (begin, end, text) tuples.
    """
    segments = []
    seg_start = 0
    n = len(line_word_ts)

    for i in range(n):
        seg_begin = line_word_ts[seg_start]["start"]
        seg_end = line_word_ts[i]["end"]
        seg_text = " ".join(line_words[seg_start:i + 1])
        too_long_time = (seg_end - seg_begin) > max_secs
        too_long_chars = len(seg_text) > max_chars
        # If word i pushes the segment past a limit and we already have at least
        # one earlier word, close the segment *before* word i.
        if (too_long_time or too_long_chars) and i > seg_start:
            segments.append((
                line_word_ts[seg_start]["start"],
                line_word_ts[i - 1]["end"],
                " ".join(line_words[seg_start:i]),
            ))
            seg_start = i

    if seg_start < n:
        segments.append((
            line_word_ts[seg_start]["start"],
            line_word_ts[n - 1]["end"],
            " ".join(line_words[seg_start:]),
        ))

    return segments


def parse_script(config):
    """Extract script lines per the config (docx or inline), emitting
    section_match messages. Returns (matched_heading_or_None, lines)."""
    docx_path = config.get("docx_path")
    video_stem = config.get("video_stem")
    filter_patterns = config.get("filter_patterns", [r"^\[.*\]$"])
    strip_patterns = config.get("strip_patterns", [])

    if not docx_path:
        return None, config.get("script_lines", [])

    if not os.path.isfile(docx_path):
        error(f"Script file not found: {docx_path}")

    log(f"parsing docx: {os.path.basename(docx_path)} "
        f"(video_stem={video_stem!r}, {len(filter_patterns)} filters, {len(strip_patterns)} strips)")
    try:
        # Try section-based extraction if we have a video stem
        if video_stem:
            matched_heading, section_lines = extract_section_lines(
                docx_path, video_stem, filter_patterns=filter_patterns, strip_patterns=strip_patterns
            )
            if section_lines is not None:
                log(f"matched section: {matched_heading!r} ({len(section_lines)} lines)")
                section_match(True, matched_heading)
                return matched_heading, section_lines
            log("no section matched — using whole document")
            section_match(False, None)
        return None, extract_script_lines(docx_path, filter_patterns=filter_patterns, strip_patterns=strip_patterns)
    except SystemExit:
        raise
    except Exception as e:
        error(f"Failed to parse .docx: {e}")


def main():
    # Read input from stdin
    try:
        raw = sys.stdin.read()
        config = json.loads(raw)
    except Exception as e:
        error(f"Failed to parse input: {e}")

    log_environment()

    mode = config.get("mode", "align")
    log(f"mode: {mode}")

    # Extract-only mode: parse the .docx and return the lines — no alignment.
    if mode == "extract":
        heading, script_lines = parse_script(config)
        extract_result(heading, script_lines)
        log(f"extract complete: {len(script_lines)} lines")
        return

    audio_path = config.get("audio_path")
    output_path = config.get("output_path")
    model_path = config.get("model_path")

    if not audio_path or not output_path:
        error("Missing audio_path or output_path")

    _, script_lines = parse_script(config)

    if not script_lines:
        error("No spoken lines found in the script")

    if not os.path.isfile(audio_path):
        error(f"Audio file not found: {audio_path}")

    total_script_words = sum(len(line.split()) for line in script_lines)
    log(f"script: {len(script_lines)} lines, {total_script_words} words")

    # Set model path for ctc_forced_aligner if provided
    if model_path:
        model_dir = os.path.dirname(model_path)
        os.makedirs(model_dir, exist_ok=True)
        # The AlignmentSingleton looks for the model at ~/ctc_forced_aligner/model.onnx
        # We symlink our model location if needed
        default_model_dir = os.path.expanduser("~/ctc_forced_aligner")
        default_model_path = os.path.join(default_model_dir, "model.onnx")
        if os.path.isfile(model_path) and not os.path.isfile(default_model_path):
            os.makedirs(default_model_dir, exist_ok=True)
            try:
                os.symlink(model_path, default_model_path)
                log(f"symlinked model: {model_path} -> {default_model_path}")
            except OSError as e:
                # If this fails AND no model exists at the default path, the
                # aligner library may silently re-download ~1.2 GB.
                log(f"WARNING: model symlink failed ({e}); "
                    f"aligner may download its own copy to {default_model_dir}")

    # Import alignment libraries
    progress("Importing alignment libraries...", 5)
    log("importing ctc_forced_aligner...")
    try:
        from ctc_forced_aligner import (
            AlignmentSingleton,
            generate_emissions,
            get_alignments,
            get_spans,
            load_audio,
            postprocess_results,
            preprocess_text,
        )
    except ImportError as e:
        error(f"Failed to import ctc_forced_aligner: {e}")

    # Run alignment pipeline
    full_text = " ".join(script_lines)

    progress("Loading alignment model...", 10)
    log("loading alignment model...")
    t0 = time.time()
    try:
        aligner = AlignmentSingleton()
    except Exception as e:
        error(f"Failed to load alignment model: {e}")
    log(f"model loaded in {time.time() - t0:.1f}s")

    progress("Loading audio...", 20)
    try:
        audio_waveform = load_audio(audio_path)
    except Exception as e:
        error(f"Failed to load audio: {e}")

    audio_seconds = len(audio_waveform) / 16000.0
    log(f"audio: {audio_seconds:.1f}s at 16kHz")
    info(total_script_words, audio_seconds)

    # Emissions (model inference over the whole audio) dominate the run time
    # and the library offers no progress callback, so a heartbeat thread
    # estimates progress from elapsed time: asymptotic toward 50%, never past.
    progress("Transcribing audio (model inference)...", 30)
    log("generating emissions...")
    t0 = time.time()
    stop_heartbeat = threading.Event()

    def heartbeat():
        # Time constant scaled to audio length: longer audio = slower ramp.
        # Calibrated against Apple-silicon runs (~0.13s inference per 1s of
        # audio) so the bar sits near the top of the band as inference ends.
        tau = max(5.0, audio_seconds * 0.12)
        while not stop_heartbeat.wait(2.0):
            elapsed = time.time() - t0
            est = 30 + 20 * (1 - math.exp(-elapsed / tau))
            progress("Transcribing audio (model inference)...", round(min(est, 49.5), 1))

    heartbeat_thread = threading.Thread(target=heartbeat, daemon=True)
    heartbeat_thread.start()
    try:
        emissions, stride = generate_emissions(
            aligner.alignment_model, audio_waveform, batch_size=16
        )
    except Exception as e:
        stop_heartbeat.set()
        error(f"Failed to generate emissions: {e}")
    finally:
        stop_heartbeat.set()
        heartbeat_thread.join(timeout=3)
    log(f"emissions generated in {time.time() - t0:.1f}s")

    progress("Preprocessing text...", 50)
    try:
        tokens_starred, text_starred = preprocess_text(
            full_text, romanize=True, language="eng"
        )
    except Exception as e:
        error(f"Failed to preprocess text: {e}")

    progress("Running CTC forced alignment...", 60)
    try:
        segments, scores, blank_token = get_alignments(
            emissions, tokens_starred, aligner.alignment_tokenizer
        )
    except Exception as e:
        error(f"Failed to run alignment: {e}")

    progress("Post-processing results...", 80)
    try:
        spans = get_spans(tokens_starred, segments, blank_token)
        word_timestamps = postprocess_results(text_starred, spans, stride, scores)
    except Exception as e:
        error(f"Failed to post-process: {e}")

    # The line-to-timestamp mapping below assumes one timestamp per
    # whitespace-separated word; if tokenization disagreed, timing drifts.
    if len(word_timestamps) != total_script_words:
        log(f"WARNING: word count mismatch — script has {total_script_words} words "
            f"but aligner returned {len(word_timestamps)} timestamps; "
            f"cue timing may drift toward the end of the video")

    progress("Generating SRT file...", 90)

    # Map word-level timestamps back to original script lines.
    # DCMP / CEA-608 limits: ≤6s on screen and ≤64 chars per cue (2×32-char
    # lines). The Swift compliance pass inserts the actual line break later.
    MAX_SUBTITLE_SECS = 6.0
    MAX_CHARS_PER_CUE = 64
    srt_entries = []
    word_idx = 0
    total_words = len(word_timestamps)

    for line in script_lines:
        line_words = line.split()
        n_words = len(line_words)

        if word_idx >= total_words:
            break

        end_idx = min(word_idx + n_words, total_words)
        line_word_ts = word_timestamps[word_idx:end_idx]

        if line_word_ts:
            begin = line_word_ts[0]["start"]
            end = line_word_ts[-1]["end"]
            duration = end - begin

            # Split if the line is too long either in time or in characters.
            if duration <= MAX_SUBTITLE_SECS and len(line) <= MAX_CHARS_PER_CUE:
                srt_entries.append((begin, end, line))
            else:
                srt_entries.extend(_split_line_segments(
                    line_words, line_word_ts, MAX_SUBTITLE_SECS, MAX_CHARS_PER_CUE
                ))

        word_idx = end_idx

    # Enforce minimum subtitle duration (DCMP minimum on-screen time)
    MIN_SUBTITLE_SECS = 1.3
    for i in range(len(srt_entries)):
        begin, end, text = srt_entries[i]
        if end - begin < MIN_SUBTITLE_SECS:
            new_end = begin + MIN_SUBTITLE_SECS
            # Don't overlap into the next cue
            if i + 1 < len(srt_entries):
                next_begin = srt_entries[i + 1][0]
                new_end = min(new_end, next_begin)
            srt_entries[i] = (begin, new_end, text)

    # Write SRT file
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            for i, (begin, end, text) in enumerate(srt_entries, 1):
                f.write(f"{i}\n")
                f.write(f"{format_srt_time(begin)} --> {format_srt_time(end)}\n")
                f.write(f"{text}\n")
                f.write("\n")
    except Exception as e:
        error(f"Failed to write SRT file: {e}")

    progress("Complete", 100)
    log(f"wrote {len(srt_entries)} cues to {output_path}")
    result(output_path, len(srt_entries))


if __name__ == "__main__":
    main()
